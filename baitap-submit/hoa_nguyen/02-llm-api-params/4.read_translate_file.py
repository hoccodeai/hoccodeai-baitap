from docx import Document
from openai import OpenAI



 # Init openAI client
client = OpenAI(
    base_url="https://api.together.xyz/v1",
    # Làm theo hướng dẫn trong bài, truy cập https://api.together.ai/settings/api-keys để lấy API Key nha
    api_key='010dde7e500cc647a52d1a92efd2c9ece8331a059942d6151cf916a4cb206420',
)


def translate_text(text, from_lan='en', to_lang='vi'):

    # gui yeu cau va nhan ket qua tu openAi
    messages = []
    messages.append({'role':'system','content':'Bạn là một thông dịch viên'})
    messages.append({'role' : 'user', 'content': f'hãy dich cho tôi từ tiếng {from_lan} sang tiếng {to_lang} cho đoạn text sau : \n\n{text} '})
    response = client.chat.completions.create(
        model = "meta-llama/Llama-3-70b-chat-hf",
        messages=messages,
        max_tokens=1000,
        temperature=0.5
    )
    bot_reply = response.choices[0].message.content
    return bot_reply


def save_file(content, file_path):
    output_doc = Document()
    for paragraph in content:
        output_doc.add_paragraph(paragraph)
    output_doc.save(file_path)
    print(f'da luu vao file {file_path}')
    

def translate_file(file_path):
    
    doc_content = Document(file_path)
    paragraphs = []    
    # duyet qua cac doan van ban trong file
    for i, para in enumerate(doc_content.paragraphs) :
      paragraph_translated = translate_text( para.text.strip())
      paragraphs.append(paragraph_translated)

    print(paragraphs)
    save_file(paragraphs,"out_file.docx")

if __name__ == '__main__':
    translate_file('test_file.docx')
 